#[Headder]
import os
import requests
import json
import docx
from docx import Document, document
from tqdm import tqdm
import subprocess

#[Define]
##define variable
input_path = "input" 
output_path = "output" 
total_paragraphs = ""
LLm_model = []
##define api
OLLAMA_API_URL = "http://localhost:11434/api/generate"
##define function
def line_counter(input_doc,head_style,head_content):
    for i, paragraph in enumerate(input_doc.paragraphs): # Use enumerate for index and value
        if paragraph.style.name == head_style and head_content in paragraph.text:  
            return i
            # Stop searching once the heading is found
            break    
def docx_name_getter(input_path):
    docx_files = [os.path.basename(file) for file in os.listdir(input_path) if file.endswith(".docx")] 
    if docx_files:
        return ', '.join(docx_files)  # Join filenames with commas and spaces
    else:
        return 'idk why script aint get correct name,so i rename it like this' 
def docx_path_outputer(input_path):
    #find all .docx files in the input folder
     input_filepath = [os.path.join(input_path, file) for file in os.listdir(input_path) if file.endswith(".docx")] 
     if input_filepath:
       for file_path in input_filepath:
           return file_path
def ollama_translate(text):
    #use to tell llm what they should do
    # construct json payload
    payload = {
        "model": LLm_model, 
        "prompt": f"[you will ONLY accurately translate the following text into Chinese, and you will not say ANYTHING UNNECESSARY ]:\n{text}",
    }
    
    # use requests to post
    with requests.post(OLLAMA_API_URL, json=payload, stream=True) as response:
        response_text = ""
        for line in response.iter_lines():
            if line:
                # decode each line from json to python dit
                decoded_line = json.loads(line.decode('utf-8'))
                # extract response
                if "response" in decoded_line:
                    response_text += decoded_line['response']
                # finshed when done is true
                if decoded_line.get("done", False):
                    break
        # return full translation
        return response_text.strip()
def progress_bar_counter(input_path):
    ###input the doc file
    file_path = docx_path_outputer(input_path)
    doc = Document(file_path)
    ###make varrible int
    head_possition = 0
    end_possition = 0
    ###find both possition
    head_possition = line_counter(doc,'Heading 2','$$Content')
    end_possition = line_counter(doc,'Heading 2','$$Annotation')
    ###caculate and error checking
    i = end_possition - head_possition - 1
    return i
    if i <= 0:
        print("\033[31m[Progress bar error]\033[0m progress bar might be inaccurate due to the fromat of file") 
def ollama_list():
 # Execute the "ollama list" command and capture its output and wake ollama
 #subprocess.run(["ollama", "serve"], capture_output=True, text=True)
 result = subprocess.run(["ollama", "list"], capture_output=True, text=True)
# Parse the output and extract model information
 models = []
 for line in result.stdout.splitlines():
    if "NAME" in line:  # Skip header line
        continue
    parts = line.split()
    name = parts[0]
    id_ = parts[1]
    size = parts[2]
    modified = parts[3]
    models.append({
        "name": name,
        "id": id_,
        "size": size,
        "modified": modified
    })
# Print the available models and ask for user selection
 for i, model in enumerate(models):
    print(f"{i+1}. {model['name']}")
 choice = int(input("\033[34m[Enter mode number]\033[0m")) - 1  # Adjust for zero-based indexing
 selected_model = models[choice]
 return selected_model['name']
##app function
def inout_folder_exists(input_path,output_path):
    #use to confirm&in-not-redo inout folder
    ##output folder
    if not os.path.exists(output_path):
        os.makedirs(output_path) 
    ##input folder
    if not os.path.exists(input_path):
        os.makedirs(input_path) 
def docx_files_indicator(input_path):
    #tell u what we have found 
  docx_files = docx_path_outputer(input_path)
  if docx_files:
       print(f"\033[32m[Docx files found]\033[0m {docx_files}")
  else:
        print("\033[31m[No docx files found]\033[0m")
        exit()
def process_document(input_path, output_path):
    ###find the file path
    file_path = docx_path_outputer(input_path)  
    ###define 2 varrible to working on
    doc = Document(file_path) 
    new_doc = Document()
    ##define body style
    body_font = doc.styles['Normal'].font
    body_font.name = 'Arial'              
    body_font.size = docx.shared.Pt(10.5) 
    ###count nedded line
    heading_index = line_counter(doc,'Heading 2','$$Content')
    ###actually translate every line and put them above the original text
    for para in tqdm(doc.paragraphs[heading_index + 1:], total=total_paragraphs, desc="\033[34m[Prossesing status]\033[0m"):
        if para.style.name.startswith('Heading'):
            break  
        if para.text.strip():
            translated_text = ollama_translate(para.text) 
            new_doc.add_paragraph(translated_text)
            new_doc.add_paragraph(para.text)
    ###get the name for output file
    name= docx_name_getter(input_path)
    #add other content together in to new file
    


    #format in to general fromat


    #save file
    new_doc.save(output_path+"/"+name)
    thepath=docx_path_outputer(output_path)
    print(f"\033[32m[Document saved]\033[0m {thepath}")  


#[Application]
##ensure in&out folder exists
inout_folder_exists(input_path,output_path)
##indicate the file that have been found
docx_files_indicator(input_path)
##count for progress bar
total_paragraphs = progress_bar_counter(input_path)
##choose model
LLm_model = ollama_list()
##process files
process_document(input_path,output_path)