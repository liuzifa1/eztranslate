#[Headder]
import os
import requests
import json
import docx
from docx import Document, document
from tqdm import tqdm
import subprocess
import whisper
import librosa
import time
import threading

#[Define]
##define variable
input_path = "input" 
output_path = "output" 
##define api
OLLAMA_API_URL = "http://localhost:11434/api/generate"
##define function
###progressbar/progress animation related function
def line_finder(input_doc,head_style,head_content):
    for i, paragraph in enumerate(input_doc.paragraphs): # Use enumerate for index and value
        if paragraph.style.name == head_style and head_content in paragraph.text:  
            return i
            # Stop searching once the heading is found
            break    
def loading_animation(state=False): 
    # value "state" tells animation to start or stop by alter loading_animation_state value
    global loading_animation_state
    loading_animation_state = state
    # define a animation function
    def animation_function_block():
        # loading global varribles
        # avoid been overlaped with previous function output
        time.sleep(0.5) 
        symbols = ['-', '\\', '|', '/']
        index = 0
        while loading_animation_state:
            print(f"\r\033[34mProcessing...\033[0m{symbols[index]}", end="")
            time.sleep(0.2)
            index = (index + 1) % len(symbols)            
    # create process that holds animation
    animation_thread = threading.Thread(target=animation_function_block)
    animation_thread.start()
def function_timer(start=False):
  global startTime
  if start:
    startTime = time.time()
  else:
    elapsedTime = time.time() - startTime
    return elapsedTime
###incharge of file path&name
def mp3_name_getter(input_path):
    mp3_files = [os.path.basename(file) for file in os.listdir(input_path) if file.endswith(".mp3")] 
    if mp3_files:
        return ', '.join(mp3_files)  # Join filenames with commas and spaces
    else:
        return 'idk why script aint get correct name,so i rename it like this' 
def mp3_path_outputer(input_path):
    #find all .mp3 files in the input folder
     input_filepath = [os.path.join(input_path, file) for file in os.listdir(input_path) if file.endswith(".mp3")] 
     if input_filepath:
       for file_path in input_filepath:
           return file_path
def docx_path_outputer(input_path):
    #find all .mp3 files in the input folder
     input_filepath = [os.path.join(input_path, file) for file in os.listdir(input_path) if file.endswith(".docx")] 
     if input_filepath:
       for file_path in input_filepath:
           return file_path
### Whisper related function
def whisper_horsepower_selector():  
    models =[
    {'name': 'tiny','parameters':'39M','.en':'yep','vram':'~1GB','relative speed':'~32x'},
    {'name': 'base','parameters':'74M','.en':'yep','vram':'~1GB','relative speed':'~16x'},
    {'name': 'small','parameters':'244M','.en':'yep','vram':'~2GB','relative speed':'~6x'},            
    {'name': 'medium','parameters':'769M','.en':'yep','vram':'~5GB','relative speed':'~2x'},
    {'name': 'large','parameters':'1550M','.en':'nope','vram':'~10GB','relative speed':'1x'}
        ]
    # Print the available models and ask for user selection
    print("\033[34m[Whisper]\033[0m Choose Whisper model density")
    print("DENSITY        VRAM     SPEED")
    for i, model in enumerate(models):
        print(f"{i+1}. {model['name']}       {model['vram']}      {model['relative speed']}")
    choice = int(input("\033[34m[Enter model number]\033[0m: ")) - 1  # Adjust for zero-based indexing
    selected_model = models[choice]['name']
    return selected_model
###ollama related function
def ollama_translate(text):
    #use to tell llm what they should do
    # construct json payload
    payload = {
        "model": LLm_model, 
        "prompt": f"[you will ONLY accurately translate the following text into Chinese, and you will not say ANYTHING UNNECESSARY ]:\n{text}",
    }
    
    # use requests to post
    with requests.post(OLLAMA_API_URL, json=payload, stream=True) as response:
        response_text = ""
        for line in response.iter_lines():
            if line:
                # decode each line from json to python dit
                decoded_line = json.loads(line.decode('utf-8'))
                # extract response
                if "response" in decoded_line:
                    response_text += decoded_line['response']
                # finshed when done is true
                if decoded_line.get("done", False):
                    break
        # return full translation
        return response_text.strip()
def ollama_list():
 # wake ollama
 # subprocess.run(["ollama","serve"])
 #list model
 result = subprocess.run(["ollama", "list"], capture_output=True, text=True)
# Parse the output and extract model information
 models = []
 for line in result.stdout.splitlines():
    if "NAME" in line:  # Skip header line
        continue
    parts = line.split()
    name = parts[0]
    id_ = parts[1]
    size = parts[2]
    modified = parts[3]
    models.append({
        "name": name,
        "id": id_,
        "size": size,
        "modified": modified
    })
# Print the available models and ask for user selection
 print("\033[34m[Ollama]\033[0m Choose working llm")
 print("    NAME")
 for i, model in enumerate(models):
    print(f"{i+1}. {model['name']}")
 choice = int(input("\033[34m[Enter model number]\033[0m: ")) - 1  # Adjust for zero-based indexing
 selected_model = models[choice]
 return selected_model['name']
##app function
def inout_folder_exists(input_path,output_path):
    #use to confirm&in-not-redo inout folder
    ##output folder
    if not os.path.exists(output_path):
        os.makedirs(output_path) 
    ##input folder
    if not os.path.exists(input_path):
        os.makedirs(input_path) 
def audio_files_indicator(input_path):
      #tell u what we have found 
  audio_files = mp3_path_outputer(input_path)
  if audio_files:
       print(f"\033[32m[Audio files found]\033[0m {audio_files}")
  else:
        print("\033[34m[No audio files found]\033[0m")
        exit()
def process_audio(input_path):
    # let u choose the Whisper density
    model_density = whisper_horsepower_selector()
    # prepare for audio processsing
    # get file & tiny model for language identify
    mp3_file= mp3_path_outputer(input_path)
    model = whisper.load_model("tiny")
    # load audio and pad/trim it to fit 30 seconds
    audio = whisper.load_audio(mp3_file)
    identify_audio = whisper.pad_or_trim(audio)
    # make log-Mel spectrogram and move to the same device as the model
    mel = whisper.log_mel_spectrogram(identify_audio).to(model.device)
    # detect the spoken language
    _, probs = model.detect_language(mel)
    print(f"\033[34m[Detected language]\033[0m {max(probs, key=probs.get)}")
    # hand file directly to large model or identiy if its nedded to be handed to en specific
    if model_density == "large":
        print("\033[34m[Model didn't switched]\033[0m large model dosen't have that")
    else:
        if max(probs, key=probs.get) == "en":
            model = whisper.load_model(f"{model_density}.en") 
            print(f"\033[34m[Model switched]\033[0m <{model_density}.en>")
        else:
            model = whisper.load_model(model_density)

    #main process
    loading_animation(True)
    function_timer(True)
    result = model.transcribe(mp3_file)
    elapsed_time = function_timer(False)
    loading_animation(False)
    print(f"\r\033[32m[Transcribe finished]\033[0m Whisper in ur computer for like {elapsed_time:.0f} seconds, then ur computer giving that up")
    return result
def process_document(input_path,output_path,input_file):
    #preparing docx
    doc = Document()
    new_doc = Document()
    doc.add_paragraph("$$Content",style='Heading2')
    doc.add_paragraph("")
    for segment in input_file['segments']:
        text = segment['text']
        doc.add_paragraph(text,style='Normal')
    doc.add_paragraph("$$Annoation",style='Heading2')    
    ###count nedded line&find head&find end&count number
    head_num = line_finder(doc,'Heading 2','$$Content')
    ass_num=line_finder(doc,'Heading 2','$$Annoation')
    total_paragraphs = ass_num - head_num
    ###actually translate every line and put them above the original text
    for para in tqdm(doc.paragraphs[head_num + 1:], total=total_paragraphs, desc="\033[34m[Prossesing status]\033[0m"):
        if para.style.name.startswith('Heading'):
            break  
        if para.text.strip():
            translated_text = ollama_translate(para.text) 
            new_doc.add_paragraph(translated_text)
            new_doc.add_paragraph(para.text)
    ###get the name for output file
    name= mp3_name_getter(input_path)
    #save file
    new_doc.save(output_path+"/"+name+".docx")
    the_path=docx_path_outputer(output_path)
    print(f"\033[32m[Document saved]\033[0m {the_path}")  


#[Application]
##ensure in&out folder exists
inout_folder_exists(input_path,output_path)
##indicate the file that have been found
audio_files_indicator(input_path)
#use whisper to process audio&skip if file dosen't exist
text_original=process_audio(input_path)
##choose model
LLm_model = ollama_list()
##process files
process_document(input_path,output_path,text_original)